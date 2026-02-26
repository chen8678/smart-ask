"""
DOCX文件处理工具
"""
import logging
from typing import Optional
from docx import Document
from io import BytesIO

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """DOCX文件处理器"""
    
    @staticmethod
    def is_docx_file(content: bytes) -> bool:
        """
        检查是否为DOCX文件
        
        Args:
            content: 文件字节内容
            
        Returns:
            是否为DOCX文件
        """
        # DOCX文件以PK开头（ZIP格式）
        return content.startswith(b'PK\x03\x04')
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Optional[str]:
        """
        从DOCX文件中提取文本内容
        
        Args:
            file_content: DOCX文件的字节内容
            
        Returns:
            提取的文本内容，如果失败返回None
        """
        try:
            # 使用python-docx库解析DOCX文件
            doc = Document(BytesIO(file_content))
            
            # 提取所有段落的文本
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            # 合并所有文本
            full_text = '\n'.join(text_content)
            
            if full_text.strip():
                logger.info(f"成功提取DOCX文本，长度: {len(full_text)} 字符")
                return full_text
            else:
                logger.warning("DOCX文件中没有找到文本内容")
                return None
                
        except Exception as e:
            logger.error(f"提取DOCX文本失败: {str(e)}")
            return None
    
    @staticmethod
    def get_docx_info(file_content: bytes) -> dict:
        """
        获取DOCX文件信息
        
        Args:
            file_content: DOCX文件的字节内容
            
        Returns:
            文件信息字典
        """
        try:
            doc = Document(BytesIO(file_content))
            
            # 获取文档属性
            core_props = doc.core_properties
            
            info = {
                'title': core_props.title or '未知标题',
                'author': core_props.author or '未知作者',
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'page_count': len(doc.paragraphs) // 20  # 估算页数
            }
            
            return info
            
        except Exception as e:
            logger.error(f"获取DOCX信息失败: {str(e)}")
            return {
                'title': '未知标题',
                'author': '未知作者',
                'created': None,
                'modified': None,
                'paragraph_count': 0,
                'table_count': 0,
                'page_count': 0
            }
